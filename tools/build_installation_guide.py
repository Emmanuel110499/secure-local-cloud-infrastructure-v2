from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "INSTALLATION.md"
OUT = ROOT / "output" / "release-v1.0.0"
OUT.mkdir(parents=True, exist_ok=True)

NAVY = "0B1F3A"
BLUE = "2563EB"
PALE = "E8EEF5"
INK = "14213D"
MUTED = "5F6F89"


def parse_md(text):
    lines = text.splitlines()
    blocks, i = [], 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            lang, buf = line[3:].strip(), []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            blocks.append(("code", lang, "\n".join(buf)))
        elif line.startswith("| "):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", c or "-") for c in cells): rows.append(cells)
                i += 1
            blocks.append(("table", rows)); i -= 1
        elif line.startswith("# "): blocks.append(("title", line[2:]))
        elif line.startswith("## "): blocks.append(("h1", line[3:]))
        elif line.startswith("### "): blocks.append(("h2", line[4:]))
        elif re.match(r"^[-*] ", line): blocks.append(("bullet", line[2:]))
        elif re.match(r"^\d+\. ", line): blocks.append(("number", re.sub(r"^\d+\. ", "", line)))
        elif line.startswith("> "): blocks.append(("callout", line[2:]))
        elif line.strip(): blocks.append(("p", line.strip()))
        i += 1
    return blocks


def clean(s):
    return re.sub(r"`([^`]+)`", r"\1", s).replace("**", "")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def build_docx(blocks, path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.65)
    sec.left_margin = Inches(.82); sec.right_margin = Inches(.82)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.2
    for name, size, color, before, after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,"1F4D78",10,5)]:
        st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    header = sec.header.paragraphs[0]
    header.text = "SECURE LOCAL CLOUD  |  GUIDE D'INSTALLATION v1.0.0"
    header.runs[0].font.name="Calibri"; header.runs[0].font.size=Pt(8); header.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
    footer = sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Document public — secrets exclus").font.size=Pt(8)
    for kind, *payload in blocks:
        if kind == "title":
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(85); p.paragraph_format.space_after=Pt(14)
            r=p.add_run(clean(payload[0])); r.bold=True; r.font.name="Calibri"; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY)
            s=doc.add_paragraph("Secure Local Cloud Infrastructure v1.0.0", style=None); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
            s.runs[0].font.size=Pt(15); s.runs[0].font.color.rgb=RGBColor.from_string(BLUE)
            doc.add_paragraph("De la création des machines virtuelles à la publication HTTPS, aux alertes et aux sauvegardes.").alignment=WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
        elif kind == "h1": doc.add_heading(clean(payload[0]), level=1)
        elif kind == "h2": doc.add_heading(clean(payload[0]), level=2)
        elif kind == "bullet": doc.add_paragraph(clean(payload[0]), style="List Bullet")
        elif kind == "number": doc.add_paragraph(clean(payload[0]), style="List Number")
        elif kind == "callout":
            p=doc.add_paragraph(); shade_cell if False else None
            p.paragraph_format.left_indent=Inches(.15); p.paragraph_format.right_indent=Inches(.15)
            r=p.add_run(clean(payload[0])); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
        elif kind == "code":
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.12); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
            for n,line in enumerate(payload[1].splitlines()):
                r=p.add_run(line + ("\n" if n < len(payload[1].splitlines())-1 else "")); r.font.name="Consolas"; r.font.size=Pt(8); r.font.color.rgb=RGBColor(255,255,255)
            pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"), NAVY); pPr.append(shd)
        elif kind == "table":
            rows=payload[0]
            if not rows: continue
            t=doc.add_table(rows=len(rows), cols=max(map(len,rows))); t.autofit=False; t.style="Table Grid"
            widths=[Inches(6.5/max(map(len,rows)))]*max(map(len,rows))
            for ri,row in enumerate(rows):
                for ci,val in enumerate(row):
                    c=t.cell(ri,ci); c.width=widths[ci]; c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    c.text=clean(val)
                    for p in c.paragraphs:
                        p.paragraph_format.space_after=Pt(2)
                        for r in p.runs: r.font.size=Pt(8.5); r.bold=(ri==0)
                    if ri==0: shade_cell(c, PALE)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
        else: doc.add_paragraph(clean(payload[0]))
    doc.save(path)


def build_pdf(blocks, path):
    base=getSampleStyleSheet()
    ss={
        "title":ParagraphStyle("title",parent=base["Title"],fontName="Helvetica-Bold",fontSize=26,leading=31,textColor=colors.HexColor("#0B1F3A"),alignment=TA_CENTER,spaceAfter=14),
        "h1":ParagraphStyle("h1",parent=base["Heading1"],fontName="Helvetica-Bold",fontSize=16,leading=20,textColor=colors.HexColor("#2563EB"),spaceBefore=12,spaceAfter=7),
        "h2":ParagraphStyle("h2",parent=base["Heading2"],fontName="Helvetica-Bold",fontSize=12,leading=15,textColor=colors.HexColor("#1F4D78"),spaceBefore=9,spaceAfter=5),
        "p":ParagraphStyle("p",parent=base["BodyText"],fontName="Helvetica",fontSize=8.7,leading=12.2,textColor=colors.HexColor("#14213D"),spaceAfter=5),
        "bullet":ParagraphStyle("bullet",parent=base["BodyText"],fontSize=8.7,leading=12,leftIndent=12,firstLineIndent=-6,spaceAfter=3),
        "callout":ParagraphStyle("callout",parent=base["BodyText"],fontSize=9,leading=13,textColor=colors.HexColor("#174EA6"),backColor=colors.HexColor("#EAF2FF"),borderPadding=8,spaceAfter=7),
        "code":ParagraphStyle(
            "code",
            parent=base["Code"],
            fontName="Courier",
            fontSize=7.7,
            leading=9.6,
            textColor=colors.HexColor("#10233F"),
            backColor=colors.HexColor("#EAF1FA"),
            borderColor=colors.HexColor("#B8CBE3"),
            borderWidth=0.6,
            borderPadding=8,
            spaceBefore=3,
            spaceAfter=9,
        ),
    }
    def esc(s): return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
    story=[]
    for kind,*payload in blocks:
        if kind=="title":
            story += [Spacer(1,45*mm),Paragraph(esc(clean(payload[0])),ss["title"]),Paragraph("Secure Local Cloud Infrastructure v1.0.0",ParagraphStyle("sub",parent=ss["p"],fontSize=14,textColor=colors.HexColor("#2563EB"),alignment=TA_CENTER)),Spacer(1,10*mm),Paragraph("Guide reproductible : VM, réseaux, sécurité, application, observabilité, alertes, sauvegardes et restauration.",ParagraphStyle("lead",parent=ss["p"],fontSize=11,leading=16,alignment=TA_CENTER)),PageBreak()]
        elif kind in ("h1","h2"): story.append(Paragraph(esc(clean(payload[0])),ss[kind]))
        elif kind=="bullet": story.append(Paragraph("• "+esc(clean(payload[0])),ss["bullet"]))
        elif kind=="number": story.append(Paragraph("• "+esc(clean(payload[0])),ss["bullet"]))
        elif kind=="callout": story.append(Paragraph(esc(clean(payload[0])),ss["callout"]))
        elif kind=="code": story.append(Preformatted(payload[1],ss["code"],maxLineLength=105))
        elif kind=="table":
            rows=payload[0]
            if rows:
                data=[[Paragraph(esc(clean(v)),ss["p"]) for v in row] for row in rows]
                t=Table(data,repeatRows=1,colWidths=[170*mm/max(map(len,rows))]*max(map(len,rows)),hAlign="LEFT")
                t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E8EEF5")),("TEXTCOLOR",(0,0),(-1,0),colors.HexColor("#0B1F3A")),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("GRID",(0,0),(-1,-1),.4,colors.HexColor("#D8E2F0")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)])); story += [t,Spacer(1,5)]
        else: story.append(Paragraph(esc(clean(payload[0])),ss["p"]))
    def page(canvas,doc):
        canvas.saveState(); canvas.setFillColor(colors.HexColor("#0B1F3A")); canvas.rect(0,A4[1]-10*mm,A4[0],10*mm,fill=1,stroke=0); canvas.setFillColor(colors.white); canvas.setFont("Helvetica-Bold",7.5); canvas.drawString(18*mm,A4[1]-6.5*mm,"SECURE LOCAL CLOUD | INSTALLATION v1.0.0"); canvas.setFillColor(colors.HexColor("#5F6F89")); canvas.drawRightString(A4[0]-18*mm,8*mm,f"Page {doc.page}"); canvas.restoreState()
    pdf=SimpleDocTemplate(str(path),pagesize=A4,leftMargin=18*mm,rightMargin=18*mm,topMargin=18*mm,bottomMargin=15*mm,title="Guide complet d'installation Secure Local Cloud",author="Secure Local Cloud Infrastructure")
    pdf.build(story,onFirstPage=page,onLaterPages=page)


if __name__ == "__main__":
    blocks=parse_md(SOURCE.read_text(encoding="utf-8"))
    build_docx(blocks,OUT/"guide-installation-complet-secure-local-cloud-v1.0.0.docx")
    build_pdf(blocks,OUT/"guide-installation-complet-secure-local-cloud-v1.0.0.pdf")
    (OUT/"INSTALLATION.md").write_text(SOURCE.read_text(encoding="utf-8"),encoding="utf-8")
    print(OUT)
